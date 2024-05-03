from collections import Counter
from pathlib import Path
from pprint import pprint

from docx import Document
from docx.text.hyperlink import Hyperlink
from docx.text.run import Run

from utils import text_skipping


def header_footer_to_markdown(header_footer, style_to_prefix):
    lines = []
    for paragraph in header_footer.paragraphs:
        if len(paragraph.text.strip()) == 0:
            continue
        lines.append(f"{style_to_prefix[paragraph.style.name]}{paragraph.text}")

    return lines


def paragraph_to_markdown(document):
    lines = []

    for i, paragraph in enumerate(document.paragraphs):
        # lines.append(f"# PARAGRAPH {i+1}")
        for content in paragraph.iter_inner_content():
            font_sizes = []
            if isinstance(content, Run):
                font_size = content.style.font.size
                font_sizes.append(font_size)

            elif isinstance(content, Hyperlink):
                font_size = content.style.font.size
            else:
                assert False

            if len(content.text.strip()) == 0:
                continue

            lines.append(f"{content.text}")

        if len(paragraph.text.strip()) == 0:
            continue

        font_size = paragraph.style.font.size
        lines.append(paragraph.text)

    return lines


def to_markdown(input_file):
    document = Document(input_file)

    lines = []
    style_to_prefix = {
        "title": "# ",
        "body": "",
        "Attribution": "# ",
        "Subheading": "## ",
        "Title 2": "### ",
        "Body": "",
        "Heading": "# ",
        "Header & Footer": "### ",
    }

    skip_counter = Counter()

    for i, section in enumerate(document.sections):
        # lines.append(f"# SECTION {i+1}")

        lines += header_footer_to_markdown(section.header, style_to_prefix)

        for section_inner in section.iter_inner_content():
            if len(section_inner.text.strip()) == 0:
                continue

            text = section_inner.text.replace("\n", " ")

            if text_skipping.should_skip(text):
                skip_counter[text] += 1
                continue

            lines.append(f"{style_to_prefix[section_inner.style.name]}{text}")

        lines += header_footer_to_markdown(section.footer, style_to_prefix)

    should_use_paragraphs = False

    if should_use_paragraphs:
        lines += paragraph_to_markdown(document)

    pprint(skip_counter)

    return "\n\n".join(lines)


def _main():
    data_dir = Path(__file__).parents[2] / "data"
    input_filename = data_dir / "sample.docx"
    output = to_markdown(input_file=input_filename)

    output_dir = (
        Path(__file__).parents[2] / "output" / "doc" / f"{input_filename.name}.md"
    )
    output_dir.parent.mkdir(parents=True, exist_ok=True)

    output_dir.write_text(output)


if __name__ == "__main__":
    _main()
